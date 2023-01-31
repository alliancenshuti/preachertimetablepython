from docx import Document
from docx.shared import Inches

document = Document()
section = document.sections[0]
header = section.header
docHeader = header.paragraphs[0]
docHeader.text = "\tNu-vision High School\n\n\tWeekly Devotional preachers List \n\tFrom Tue .... TO Thur ..... ..... 2023"
docHeader.style = document.styles["Header"]

document.add_heading(' ', 0)


def output_timetable(preachersAssigned, streams):
    table = document.add_table(rows=len(preachersAssigned) + 1, cols=4)
    # table heading
    tableHeader = table.rows[0].cells
    tableHeader[0].text = "class"
    tableHeader[1].text = "tuesday"
    tableHeader[2].text = "wednesday"
    tableHeader[3].text = "thursday"
    # set all the classes
    for i in range(1, len(preachersAssigned) + 1):
        cell = table.cell(i, 0)
        i -= 1
        cell.text = streams[i]
    for i in range(1, len(preachersAssigned) + 1):
        stream = table.rows[i].cells
        i -= 1
        stream[1].text = preachersAssigned[i]["tuesday"]
        stream[2].text = preachersAssigned[i]["wednesday"]
        stream[3].text = preachersAssigned[i]["thursday"]
    table.style = document.styles['Table Grid']
    document.add_page_break()

    document.save('timetable.docx')
